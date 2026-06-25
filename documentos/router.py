from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from typing import Optional
import pandas as pd
from docx import Document
from fpdf import FPDF

from shared.firebase import db

router = APIRouter(prefix="/exportar", tags=["documentos"])


# --- FUNCIONES DE FILTRADO NÚCLEO ---
def obtener_datos_filtrados(cuestionario_id: str, distrito: str = None, seccion: str = None, sexo: str = None):
    docs = db.collection("registros_campo").where("cuestionario_id", "==", cuestionario_id).stream()
    registros = []

    for doc in docs:
        data = doc.to_dict()
        fila = {
            "ID_Registro": doc.id,
            "Usuario": data.get("usuario_id", ""),
            "Seccion": data.get("seccion_ine", ""),
            "Fecha": data.get("fecha_captura", "")
        }

        # Extraer demográficos
        demograficos = data.get("datos_demograficos", {})
        fila.update(demograficos)

        # Extraer respuestas
        respuestas = data.get("respuestas", {})
        for preg, resp in respuestas.items():
            fila[f"P_{preg}"] = str(resp)

        registros.append(fila)

    if not registros:
        raise HTTPException(status_code=404, detail="No hay datos con esos parámetros")

    df = pd.DataFrame(registros)

    # Aplicar filtros dinámicos
    if distrito:
        pass  # Aquí puedes cruzar con tu catálogo de secciones-distritos
    if seccion:
        df = df[df["Seccion"] == seccion]
    if sexo:
        df = df[df["sexo"] == sexo]

    return df


# --- RUTAS DE EXPORTACIÓN ---

@router.get("/excel/{cuestionario_id}")
async def exportar_excel(
    cuestionario_id: str,
    seccion: Optional[str] = Query(None),
    sexo: Optional[str] = Query(None)
):
    df = obtener_datos_filtrados(cuestionario_id, None, seccion, sexo)
    file_path = f"/tmp/reporte_{cuestionario_id}.xlsx"
    df.to_excel(file_path, index=False, engine='openpyxl')
    return FileResponse(file_path, filename="Reporte_Crudo.xlsx")


@router.get("/word/{cuestionario_id}")
async def exportar_word(cuestionario_id: str, seccion: Optional[str] = Query(None)):
    df = obtener_datos_filtrados(cuestionario_id, None, seccion, None)

    doc = Document()
    doc.add_heading(f'Resumen Ejecutivo - Encuesta: {cuestionario_id}', 0)
    doc.add_paragraph(f'Total de encuestas validadas: {len(df)}')
    if seccion:
        doc.add_paragraph(f'Filtro aplicado - Sección: {seccion}')

    # Tabla resumen
    doc.add_heading('Muestra Demográfica', level=1)
    conteo_sexo = df['sexo'].value_counts() if 'sexo' in df.columns else {}
    for genero, total in conteo_sexo.items():
        doc.add_paragraph(f'{genero}: {total} encuestas', style='List Bullet')

    file_path = f"/tmp/ejecutivo_{cuestionario_id}.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename="Reporte_Ejecutivo.docx")


@router.get("/pdf/{cuestionario_id}")
async def exportar_pdf(cuestionario_id: str):
    df = obtener_datos_filtrados(cuestionario_id)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=16, style='B')
    pdf.cell(200, 10, txt="Reporte General SISAT", ln=True, align='C')

    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Total de registros: {len(df)}", ln=True)

    file_path = f"/tmp/reporte_{cuestionario_id}.pdf"
    pdf.output(file_path)
    return FileResponse(file_path, filename="Reporte_General.pdf")

import os
import json
import base64
from fastapi import FastAPI, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
import firebase_admin
from firebase_admin import credentials, auth, firestore
import pandas as pd
from docx import Document
from fpdf import FPDF
from typing import Optional
import boto3
from botocore.client import Config as BotoConfig

# --- INICIALIZACIÓN SEGURA DE FIREBASE ---
# Railway leerá el JSON de Firebase desde una variable de entorno en Base64 para evitar errores de saltos de línea
encoded_creds = os.environ.get("FIREBASE_B64")
if encoded_creds:
    decoded_creds = base64.b64decode(encoded_creds).decode('utf-8')
    cred_dict = json.loads(decoded_creds)
    cred = credentials.Certificate(cred_dict)
else:
    # Fallback para pruebas locales en tu PC
    cred = credentials.Certificate("firebase-adminsdk.json")

if not firebase_admin._apps:
    firebase_admin.initialize_app(cred)

db = firestore.client()
app = FastAPI(title="SISAT Backend", description="Motor de Exportación y Gestión ACL")

# --- CLIENTE CLOUDFLARE R2 (compatible con S3) ---
# Las credenciales viven SOLO en variables de entorno de Railway, nunca en
# el codigo ni en la app movil. Configurar en Railway:
#   R2_ENDPOINT, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY, R2_BUCKET_NAME
r2_client = boto3.client(
    "s3",
    endpoint_url=os.environ.get("R2_ENDPOINT"),
    aws_access_key_id=os.environ.get("R2_ACCESS_KEY_ID"),
    aws_secret_access_key=os.environ.get("R2_SECRET_ACCESS_KEY"),
    config=BotoConfig(signature_version="s3v4"),
    region_name="auto",
)
R2_BUCKET_NAME = os.environ.get("R2_BUCKET_NAME", "sisat-audios")

# --- FUNCIONES DE FILTRADO NÚCLEO ---
def obtener_datos_filtrados(cuestionario_id: str, distrito: str = None, seccion: str = None, sexo: str = None):
    docs = db.collection("registros_campo").where("cuestionario_id", "==", cuestionario_id).stream()
    registros = []
    
    for doc in docs:
        data = doc.to_dict()
        fila = {
            "ID_Registro": doc.id,
            "Usuario": data.get("usuario_id", ""),
            "Seccion": data.get("seccion_ine", ""),
            "Fecha": data.get("fecha_captura", "")
        }
        
        # Extraer demográficos
        demograficos = data.get("datos_demograficos", {})
        fila.update(demograficos)
        
        # Extraer respuestas
        respuestas = data.get("respuestas", {})
        for preg, resp in respuestas.items():
            fila[f"P_{preg}"] = str(resp)
            
        registros.append(fila)

    if not registros:
        raise HTTPException(status_code=404, detail="No hay datos con esos parámetros")

    df = pd.DataFrame(registros)

    # Aplicar filtros dinámicos
    if distrito:
        pass # Aquí puedes cruzar con tu catálogo de secciones-distritos
    if seccion:
        df = df[df["Seccion"] == seccion]
    if sexo:
        df = df[df["sexo"] == sexo]
        
    return df

# --- RUTAS DE EXPORTACIÓN ---

@app.get("/exportar/excel/{cuestionario_id}")
async def exportar_excel(
    cuestionario_id: str, 
    seccion: Optional[str] = Query(None), 
    sexo: Optional[str] = Query(None)
):
    df = obtener_datos_filtrados(cuestionario_id, None, seccion, sexo)
    file_path = f"/tmp/reporte_{cuestionario_id}.xlsx"
    df.to_excel(file_path, index=False, engine='openpyxl')
    return FileResponse(file_path, filename="Reporte_Crudo.xlsx")

@app.get("/exportar/word/{cuestionario_id}")
async def exportar_word(cuestionario_id: str, seccion: Optional[str] = Query(None)):
    df = obtener_datos_filtrados(cuestionario_id, None, seccion, None)
    
    doc = Document()
    doc.add_heading(f'Resumen Ejecutivo - Encuesta: {cuestionario_id}', 0)
    doc.add_paragraph(f'Total de encuestas validadas: {len(df)}')
    if seccion:
        doc.add_paragraph(f'Filtro aplicado - Sección: {seccion}')
        
    # Tabla resumen
    doc.add_heading('Muestra Demográfica', level=1)
    conteo_sexo = df['sexo'].value_counts() if 'sexo' in df.columns else {}
    for genero, total in conteo_sexo.items():
        doc.add_paragraph(f'{genero}: {total} encuestas', style='List Bullet')

    file_path = f"/tmp/ejecutivo_{cuestionario_id}.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename="Reporte_Ejecutivo.docx")

@app.get("/exportar/pdf/{cuestionario_id}")
async def exportar_pdf(cuestionario_id: str):
    df = obtener_datos_filtrados(cuestionario_id)
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=16, style='B')
    pdf.cell(200, 10, txt="Reporte General SISAT", ln=True, align='C')
    
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Total de registros: {len(df)}", ln=True)
    
    file_path = f"/tmp/reporte_{cuestionario_id}.pdf"
    pdf.output(file_path)
    return FileResponse(file_path, filename="Reporte_General.pdf")
# --- SUBIDA DE AUDIO A CLOUDFLARE R2 ---
@app.post("/audio/subir")
async def subir_audio(
    registro_id: str = Form(...),
    audio: UploadFile = File(...)
):
    """
    Recibe el audio grabado por el encuestador (desde la app, offline-first)
    y lo sube a Cloudflare R2. La app nunca toca las credenciales de R2.
    """
    if not r2_client.meta.endpoint_url or not os.environ.get("R2_ACCESS_KEY_ID"):
        raise HTTPException(
            status_code=500,
            detail="R2 no está configurado en el servidor (faltan variables de entorno)."
        )

    extension = os.path.splitext(audio.filename or "")[1] or ".m4a"
    object_key = f"audios/{registro_id}{extension}"

    try:
        contenido = await audio.read()
        r2_client.put_object(
            Bucket=R2_BUCKET_NAME,
            Key=object_key,
            Body=contenido,
            ContentType=audio.content_type or "audio/mp4",
        )
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error subiendo a R2: {e}")

    # Guarda la referencia del audio en el registro de Firestore correspondiente.
    try:
        db.collection("registros_campo").document(registro_id).update({
            "audio_key": object_key,
            "audio_sincronizado": True,
        })
    except Exception:
        # El audio ya quedó en R2 aunque el registro aun no exista en Firestore
        # (puede llegar despues, por ejemplo si la encuesta se sincroniza por separado).
        pass

    return {"status": "ok", "registro_id": registro_id, "object_key": object_key}
